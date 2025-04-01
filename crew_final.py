import time
import streamlit as st
import os
import base64
import random
import re
from pydub import AudioSegment
from openai import OpenAI
from dotenv import load_dotenv
from docx import Document
from io import BytesIO
from crewai import Agent, Task, Crew, Process
from langchain_core.language_models import BaseLLM
from langchain_core.outputs import LLMResult, Generation

load_dotenv()

# Initialize Groq client with retry logic
groq_client = OpenAI(
    api_key=os.environ["GROQ_API_KEY"],
    base_url="https://api.groq.com/openai/v1"
)

# Keep track of token usage (approximation)
if 'token_usage' not in st.session_state:
    st.session_state.token_usage = {
        'timestamp': time.time(),
        'used': 0
    }

def handle_rate_limit(max_retries=5, initial_delay=20):
    """Decorator to handle rate limit errors with exponential backoff and jitter"""
    def decorator(func):
        def wrapper(*args, **kwargs):
            retries = 0
            delay = initial_delay
            
            # Update token usage tracking
            current_time = time.time()
            if current_time - st.session_state.token_usage['timestamp'] >= 60:
                # Reset counter after a minute
                st.session_state.token_usage = {
                    'timestamp': current_time,
                    'used': 0
                }
            
            # Estimate if we're close to limit and need to wait
            if st.session_state.token_usage['used'] > 5000:  # Close to 6000 TPM limit
                wait_time = 60 - (current_time - st.session_state.token_usage['timestamp'])
                if wait_time > 0:
                    st.info(f"Approaching rate limit. Waiting {wait_time:.1f} seconds to avoid errors...")
                    time.sleep(wait_time)
                    st.session_state.token_usage = {
                        'timestamp': time.time(),
                        'used': 0
                    }
            
            while retries < max_retries:
                try:
                    result = func(*args, **kwargs)
                    
                    # Approximate token count for tracking
                    if hasattr(result, 'usage') and hasattr(result.usage, 'total_tokens'):
                        st.session_state.token_usage['used'] += result.usage.total_tokens
                    
                    return result
                except Exception as e:
                    if "rate_limit_exceeded" in str(e) or "Rate limit" in str(e):
                        retries += 1
                        if retries == max_retries:
                            st.error("Rate limit exceeded after multiple retries. Please try again later.")
                            raise
                        
                        # Add jitter to avoid thundering herd problem
                        jitter = random.uniform(0.8, 1.2)
                        actual_delay = delay * jitter
                        st.warning(f"Rate limit reached. Retrying in {actual_delay:.1f} seconds...")
                        time.sleep(actual_delay)
                        delay *= 2  # Exponential backoff
                        
                        # Reset token usage tracking after waiting
                        if actual_delay > 10:
                            st.session_state.token_usage = {
                                'timestamp': time.time(),
                                'used': 0
                            }
                    else:
                        raise
        return wrapper
    return decorator

@handle_rate_limit()
def transcribe_audio(audio_file):
    """Transcribe audio with rate limit handling"""
    return groq_client.audio.transcriptions.create(
        model="whisper-large-v3",
        file=audio_file,
        response_format="text"
    )

@handle_rate_limit()
def generate_chat_completion(messages, model="llama-3.1-8b-instant", max_tokens=150):
    """Generate chat completion with rate limit handling"""
    return groq_client.chat.completions.create(
        model=model,
        messages=messages,
        temperature=0.3,
        max_tokens=max_tokens
    )

def is_aviation_related(text):
    """Check if the transcript contains aviation-related content"""
    aviation_keywords = [
        'aircraft', 'airplane', 'helicopter', 'flight', 'pilot', 'ATC', 
        'air traffic', 'runway', 'tower', 'approach', 'landing', 'takeoff',
        'emergency', 'mayday', 'squawk', 'altitude', 'heading', 'vectors',
        'cleared', 'taxi', 'hold short', 'ILS', 'VFR', 'IFR', 'transponder'
    ]
    text_lower = text.lower()
    return any(keyword in text_lower for keyword in aviation_keywords)

def audio_to_base64(file):
    """Convert audio file to base64 for embedding in HTML"""
    with open(file, "rb") as audio_file:
        audio_bytes = audio_file.read()
        base64_audio = base64.b64encode(audio_bytes).decode()
    return base64_audio

def extract_section(text, start_marker, end_marker=None):
    """Extract a section of text between two markers"""
    if not start_marker:
        return ""
        
    start_idx = text.find(start_marker)
    if start_idx == -1:
        return ""
        
    start_idx += len(start_marker)
    if end_marker:
        end_idx = text.find(end_marker, start_idx)
        return text[start_idx:end_idx].strip() if end_idx != -1 else text[start_idx:].strip()
    return text[start_idx:].strip()

def clean_text(text):
    """Remove markdown formatting and extra spaces"""
    text = text.replace('**', '').replace('*', '').replace('#', '')
    # Collapse multiple newlines into one
    text = '\n'.join([line.strip() for line in text.split('\n') if line.strip()])
    return text

def extract_report_data(report_text):
    """Extract structured data from report text for consistent formatting"""
    data = {
        'date': extract_section(report_text, 'Date:', 'Location:').strip() or 'Not specified',
        'aircraft': extract_section(report_text, 'Aircraft:', 'Flight Number:').strip() or extract_section(report_text, 'Aircraft:', 'Location:').strip() or '',
        'location': extract_section(report_text, 'Location:', 'Aircraft:').strip() or extract_section(report_text, 'Location:', 'Incident Type:').strip() or '',
        'incident_type': extract_section(report_text, 'Incident Type:', 'Key Details:').strip() or '',
        'key_details': [],
        'safety_concerns': [],
        'recommendations': []
    }
    
    # Extract bullet points for key sections
    key_details = extract_section(report_text, 'Key Details:', 'Safety Concerns:')
    safety_concerns = extract_section(report_text, 'Safety Concerns:', 'Recommendations:')
    recommendations = extract_section(report_text, 'Recommendations:', None)
    
    # Process bullet points
    for section, target in [
        (key_details, data['key_details']),
        (safety_concerns, data['safety_concerns']),
        (recommendations, data['recommendations'])
    ]:
        for line in section.split('\n'):
            line = line.strip()
            if line:
                # Remove bullet character and add to list
                clean_line = line.lstrip('-*•').strip()
                if clean_line:
                    target.append(clean_line)
    
    return data

def generate_markdown_report(report_data):
    """Generate a properly formatted markdown report from structured data"""
    markdown = "# **Flight Incident Investigation Report**\n\n"
    
    # General Information section
    markdown += "## **General Information**\n"
    markdown += f"**Date:** {report_data.get('date', 'Not specified')}  \n"
    markdown += f"**Aircraft:** {report_data.get('aircraft', '')}  \n"
    markdown += f"**Location:** {report_data.get('location', '')}  \n"
    markdown += f"**Incident Type:** {report_data.get('incident_type', '')}\n\n"
    
    # Key Details section
    markdown += "## **Key Details**\n"
    details = report_data.get('key_details', [])
    for detail in details:
        markdown += f"* {detail}\n"
    markdown += "\n"
    
    # Safety Concerns section
    markdown += "## **Safety Concerns**\n"
    concerns = report_data.get('safety_concerns', [])
    for concern in concerns:
        markdown += f"* {concern}\n"
    markdown += "\n"
    
    # Recommendations section
    markdown += "## **Recommendations**\n"
    recommendations = report_data.get('recommendations', [])
    for recommendation in recommendations:
        markdown += f"* {recommendation}\n"
        
    return markdown

def format_report_as_html(report_text):
    """Format report into a clean HTML structure with proper headings and formatting"""
    # Extract structured data from the report
    report_data = extract_report_data(report_text)
    
    html_content = """
    <style>
        .report-container {
            font-family: Arial, sans-serif;
            line-height: 1.6;
            color: #333;
            max-width: 800px;
            margin: auto;
            padding: 20px;
            background-color: #f9f9f9;
            border-radius: 8px;
            box-shadow: 0 2px 4px rgba(0,0,0,0.1);
        }
        .report-title {
            color: #2c3e50;
            border-bottom: 2px solid #3498db;
            padding-bottom: 10px;
            margin-bottom: 20px;
            text-align: center;
            font-weight: bold;
        }
        .section {
            margin-bottom: 20px;
        }
        .section-title {
            color: #2c3e50;
            font-weight: bold;
            margin-bottom: 10px;
        }
        .section-content {
            white-space: pre-line;
            margin-left: 10px;
        }
        .section-content ul {
            margin-top: 5px;
            padding-left: 20px;
        }
        .section-content li {
            margin-bottom: 5px;
        }
    </style>
    <div class="report-container">
        <h1 class="report-title">Flight Incident Investigation Report</h1>
    """

    # Add General Information section
    html_content += '<div class="section"><h2 class="section-title">General Information</h2><div class="section-content">'
    html_content += f'<strong>Date:</strong> {report_data["date"]}<br>'
    html_content += f'<strong>Aircraft:</strong> {report_data["aircraft"]}<br>'
    html_content += f'<strong>Location:</strong> {report_data["location"]}<br>'
    html_content += f'<strong>Incident Type:</strong> {report_data["incident_type"]}<br>'
    html_content += '</div></div>'

    # Add Key Details section
    if report_data['key_details']:
        html_content += '<div class="section"><h2 class="section-title">Key Details</h2>'
        html_content += '<div class="section-content"><ul>'
        for detail in report_data['key_details']:
            html_content += f'<li>{detail}</li>'
        html_content += '</ul></div></div>'

    # Add Safety Concerns section
    if report_data['safety_concerns']:
        html_content += '<div class="section"><h2 class="section-title">Safety Concerns</h2>'
        html_content += '<div class="section-content"><ul>'
        for concern in report_data['safety_concerns']:
            html_content += f'<li>{concern}</li>'
        html_content += '</ul></div></div>'

    # Add Recommendations section
    if report_data['recommendations']:
        html_content += '<div class="section"><h2 class="section-title">Recommendations</h2>'
        html_content += '<div class="section-content"><ul>'
        for recommendation in report_data['recommendations']:
            html_content += f'<li>{recommendation}</li>'
        html_content += '</ul></div></div>'

    html_content += "</div>"
    return html_content

def generate_docx(report_data):
    """Generate a Word document from structured report data"""
    doc = Document()
    doc.add_heading('Flight Incident Investigation Report', 0)
    
    # General Information section
    doc.add_heading('General Information', level=1)
    p = doc.add_paragraph()
    p.add_run('Date: ').bold = True
    p.add_run(report_data.get('date', 'Not specified'))
    
    p = doc.add_paragraph()
    p.add_run('Aircraft: ').bold = True
    p.add_run(report_data.get('aircraft', ''))
    
    p = doc.add_paragraph()
    p.add_run('Location: ').bold = True
    p.add_run(report_data.get('location', ''))
    
    p = doc.add_paragraph()
    p.add_run('Incident Type: ').bold = True
    p.add_run(report_data.get('incident_type', ''))
    
    # Key Details section
    doc.add_heading('Key Details', level=1)
    for detail in report_data.get('key_details', []):
        doc.add_paragraph(detail, style='List Bullet')
    
    # Safety Concerns section
    doc.add_heading('Safety Concerns', level=1)
    for concern in report_data.get('safety_concerns', []):
        doc.add_paragraph(concern, style='List Bullet')
    
    # Recommendations section
    doc.add_heading('Recommendations', level=1)
    for recommendation in report_data.get('recommendations', []):
        doc.add_paragraph(recommendation, style='List Bullet')
    
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

def get_download_link(bio, filename):
    """Generate a download link for the document"""
    b64 = base64.b64encode(bio.read()).decode()
    return f'<a href="data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{b64}" download="{filename}" style="padding: 10px 20px; background-color: #3498db; color: white; text-decoration: none; border-radius: 4px; display: inline-block;">Download Report</a>'

# Custom Langchain-like LLM wrapper for Groq with rate limit handling
class GroqLLM(BaseLLM):
    """Wrapper around Groq API."""

    client: OpenAI
    model: str = "groq/llama-3.1-8b-instant"
    temperature: float = 0.7
    max_tokens: int = 1024  # Reduced from 2048 to conserve tokens
    top_p: float = 1.0

    @property
    def _llm_type(self) -> str:
        return "groq"

    @handle_rate_limit()
    def _call(self, prompt: str, stop: list[str] | None = None) -> str:
        messages = [{"role": "user", "content": prompt}]
        stop_sequences = list(stop) if isinstance(stop, tuple) else stop
        completion = self.client.chat.completions.create(
            model=self.model.split('/')[1],
            messages=messages,
            temperature=self.temperature,
            max_tokens=self.max_tokens,
            top_p=self.top_p,
            stop=stop_sequences,
        )
        return completion.choices[0].message.content

    @handle_rate_limit()
    def _generate(self, prompts: list[str], stop: list[str] | None = None) -> LLMResult:
        generations = []
        stop_sequences = list(stop) if isinstance(stop, tuple) else stop
        for prompt in prompts:
            messages = [{"role": "user", "content": prompt}]
            completion = self.client.chat.completions.create(
                model=self.model.split('/')[1],
                messages=messages,
                temperature=self.temperature,
                max_tokens=self.max_tokens,
                top_p=self.top_p,
                stop=stop_sequences,
            )
            generations.append([Generation(text=completion.choices[0].message.content)])
        return LLMResult(generations=generations)

# Streamlit UI Configuration
st.set_page_config(
    layout="wide",
    page_title="Black Box Archive Analysis"
)
st.title("Black Box Archive Analysis")

# Initialize session state for chat
if 'messages' not in st.session_state:
    st.session_state.messages = []
if 'chat_history' not in st.session_state:
    st.session_state.chat_history = []

# Create three columns (left for inputs, middle for report, right for chat)
left_col, mid_col, right_col = st.columns([1, 2, 1])

with left_col:
    st.header("Upload & Analyze")
    uploaded_file = st.file_uploader("Upload an MP3 file", type=["mp3"], key="uploader", 
                                   on_change=lambda: [st.session_state.pop(k, None) for k in ['transcript', 'report_text']])

    if uploaded_file is not None:
        with open("uploaded_file.mp3", "wb") as f:
            f.write(uploaded_file.getbuffer())

        base64_audio = audio_to_base64("uploaded_file.mp3")
        st.subheader("Your Uploaded Audio File")
        st.markdown(f"""
        <audio controls style="width: 100%;">
            <source src="data:audio/mp3;base64,{base64_audio}" type="audio/mp3">
            Your browser does not support the audio element.
        </audio>
        """, unsafe_allow_html=True)

        if st.button("Analyze", type="primary", key="analyze_btn"):
            try:
                with st.spinner("Transcribing audio..."):
                    with open("uploaded_file.mp3", "rb") as audio_file:
                        transcript = transcribe_audio(audio_file)
                
                st.session_state.transcript = transcript
                st.success("Transcription Complete!")
                st.text_area("Transcription", transcript, height=200)
                
                # CHECK IF TRANSCRIPT IS AVIATION RELATED BEFORE PROCEEDING
                with st.spinner("Validating aviation content..."):
                    is_aviation = is_aviation_related(transcript)
                
                if not is_aviation:
                    st.error("⚠️ This audio does not appear to be aviation-related.")
                    st.warning("Please upload an audio file with aviation or flight-related content.")
                    # Clear any previous report data
                    if 'report_text' in st.session_state:
                        del st.session_state.report_text
                    st.stop() 
                else:
                    st.success(f"✅ Aviation content detected!")
                    
                    with st.spinner("Generating Report using CrewAI..."):
                        # Initialize Groq LLM
                        groq_llm = GroqLLM(client=groq_client)

                        # Define Agents with more conservative settings
                        transcription_reader = Agent(
                            role="ATC Blackbox Transcription Reader",
                            goal="Read and clean ATC Blackbox transcripts",
                            backstory="Specializes in cleaning and formatting ATC transcripts",
                            verbose=True,
                            allow_delegation=False,
                            llm=groq_llm,
                            max_iter=2  # Limit processing steps
                        )

                        incident_report_writer = Agent(
                            role="Incident Reporter",
                            goal="Generate structured incident reports with the following sections: Date, Aircraft, Location, Incident Type, Key Details, Safety Concerns, and Recommendations",
                            backstory="Creates brief yet comprehensive aviation incident reports with proper formatting",
                            verbose=True,
                            allow_delegation=False,
                            llm=groq_llm,
                            max_iter=2
                        )

                        # Task for the transcription reader agent
                        clean_transcription_task = Task(
                            description=f'Clean this ATC transcript: {transcript}',
                            agent=transcription_reader,
                            expected_output='Formatted transcription as conversation.',
                        )

                        # Task for the incident report writer agent with formatted structure
                        incident_report_task = Task(
                            description="""
                            Generate a concise incident report with the following sections:
                            - Date: [date of incident]
                            - Aircraft: [aircraft type and registration]
                            - Location: [location of incident]
                            - Incident Type: [type of incident]
                            - Key Details: [bullet points of key details]
                            - Safety Concerns: [bullet points of safety concerns]
                            - Recommendations: [bullet points of recommendations]
                            
                            Use bullet points for Key Details, Safety Concerns, and Recommendations.
                            Each bullet point should start with an asterisk (*).
                            """,
                            agent=incident_report_writer,
                            expected_output='Structured incident report with all required sections.',
                            context=[clean_transcription_task]
                        )

                        # Create Crew with sequential process
                        crew = Crew(
                            agents=[transcription_reader, incident_report_writer],
                            tasks=[clean_transcription_task, incident_report_task],
                            process=Process.sequential,
                            verbose=True
                        )

                        report_text = str(crew.kickoff())
                        report_text = clean_text(report_text)
                        st.session_state.report_text = report_text

                        # Initialize chat with report context
                        st.session_state.messages = [
                            {"role": "system", "content": f"""You are an aviation assistant. Answer questions about: {report_text}"""}
                        ]

            except Exception as e:
                st.error(f"An error occurred: {str(e)}")
                if "rate_limit_exceeded" in str(e) or "Rate limit" in str(e):
                    st.info("Rate limit reached. Please wait 1-2 minutes and try again, or upgrade your Groq plan for higher limits.")
                    # Add a sleep time suggestion with a countdown
                    for i in range(60, 0, -1):
                        if i % 10 == 0:
                            st.info(f"Suggested wait time: {i} seconds remaining")
                            time.sleep(10)

with mid_col:
    st.header("Generated Report")
    if 'report_text' in st.session_state:
        # Double-check aviation content before showing report
        if 'transcript' in st.session_state:
            is_aviation= is_aviation_related(st.session_state.transcript)
            if not is_aviation:
                st.error("Content validation failed - non-aviation content detected")
                st.warning("Report generation blocked due to non-aviation content")
                del st.session_state.report_text
            else:
                # Extract structured data from the report text
                report_data = extract_report_data(st.session_state.report_text)
                
                # Display formatted report using markdown
                st.markdown(generate_markdown_report(report_data))
                
                # Generate Word document for download
                doc_bio = generate_docx(report_data)
                download_link = get_download_link(doc_bio, "flight_incident_report.docx")
                st.markdown(download_link, unsafe_allow_html=True)
    else:
        st.info("Upload an audio file and click 'Analyze' to generate a report")

with right_col:
    # Custom CSS for the compact chat interface
    st.markdown("""
    <style>
        .chat-header {
            padding: 12px;
            background: #2c3e50;
            color: white;
            font-weight: bold;
            margin: 0 !important;
        }
        .chat-content {
            flex: 1;
            display: flex;
            flex-direction: column;
            overflow: hidden;
        }
        .chat-messages {
            flex: 1;
            overflow-y: auto;
            padding: 12px;
            background: #f9f9f9;
        }
        .chat-input-container {
            padding: 12px;
            border-top: 1px solid #e1e1e1;
            background: white;
        }
        .user-message {
            background: #e3f2fd;
            border-radius: 18px 18px 0 18px;
            padding: 10px 14px;
            margin-left: auto;
            margin-bottom: 8px;
            max-width: 85%;
        }
        .assistant-message {
            background: #f5f5f5;
            border-radius: 18px 18px 18px 0;
            padding: 10px 14px;
            margin-bottom: 8px;
            max-width: 85%;
        }
        .welcome-message {
            background: #f5f5f5;
            border-radius: 12px;
            padding: 12px;
            margin-bottom: 12px;
        }
        .stTextInput>div>div>input {
            border-radius: 18px;
            padding: 8px 16px;
        }
    </style>
    """, unsafe_allow_html=True)

    # Main chat container with fixed height
    with st.container():
        # Header
        st.markdown('<div class="chat-header">Report Assistant</div>', unsafe_allow_html=True)

        # Content area (messages + input)
        st.markdown('<div class="chat-content">', unsafe_allow_html=True)

        # Messages area
        st.markdown('<div class="chat-messages">', unsafe_allow_html=True)

        # Welcome message (only first time)
        if len(st.session_state.chat_history) == 0:
            st.markdown("""
            <div class="welcome-message">
            <b>How can I help with this report?</b><br>
            Ask about:<br>
            • Incident details<br>
            • Safety recommendations<br>
            • Contributing factors
            </div>
            """, unsafe_allow_html=True)

        # Display conversation history
        for msg in st.session_state.chat_history:
            if msg['role'] == 'user':
                st.markdown(f'<div class="user-message">{msg["content"]}</div>', unsafe_allow_html=True)
            else:
                st.markdown(f'<div class="assistant-message">{msg["content"]}</div>', unsafe_allow_html=True)

        st.markdown('</div>', unsafe_allow_html=True)  # Close chat-messages

        # Input area (inside the same container)
        st.markdown('<div class="chat-input-container">', unsafe_allow_html=True)
        with st.form(key='chat_form', clear_on_submit=True):
            user_input = st.text_input(
                "Ask about the incident report",
                key="user_query",
                label_visibility="collapsed",
                placeholder="Type your question..."
            )
            submitted = st.form_submit_button("Send", type="primary")
        st.markdown('</div>', unsafe_allow_html=True)  # Close chat-input-container

        st.markdown('</div>', unsafe_allow_html=True)  # Close chat-content

    # Handle form submission
    if submitted and user_input.strip():
        if 'report_text' not in st.session_state:
            st.error("Please generate a report first")
        else:
            # Add user question to history
            st.session_state.chat_history.append({"role": "user", "content": user_input})

            try:
                prompt = f"""
                Based on this flight incident report:
                {st.session_state.report_text}

                Question: {user_input}

                Provide a concise 1-2 sentence answer.
                If the answer isn't in the report, respond: "This information is not in the report."
                """

                response = generate_chat_completion(
                    [{"role": "user", "content": prompt}],
                    max_tokens=100
                )

                answer = response.choices[0].message.content
                st.session_state.chat_history.append({"role": "assistant", "content": answer})

                # Auto-scroll to bottom
                st.markdown("""
                <script>
                    setTimeout(function() {
                        const container = window.parent.document.querySelector('.chat-messages');
                        container.scrollTop = container.scrollHeight;
                    }, 100);
                </script>
                """, unsafe_allow_html=True)

                st.rerun()

            except Exception as e:
                st.error(f"Error generating response: {str(e)}")
                if "rate_limit_exceeded" in str(e) or "Rate limit" in str(e):
                    st.info("Please wait a moment and try again.")
